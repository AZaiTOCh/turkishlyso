from __future__ import annotations

import base64
import io
import re
from pathlib import Path

from tokenish_engine.config import settings
from tokenish_engine.models import IngestResult

OCR_INTENT = re.compile(
    r"\b(read|text|ocr|transcribe|extract|scan|parse)\b",
    re.IGNORECASE,
)


def _ext(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def _slice_pages(pages: list[str], page_range: str) -> list[str]:
    m = re.match(r"^\s*(\d+)\s*-\s*(\d+)\s*$", page_range or "")
    if not m:
        return pages
    start, end = int(m.group(1)), int(m.group(2))
    start = max(1, start)
    end = min(len(pages), end)
    if start > end:
        return pages
    return pages[start - 1 : end]


def _preprocess_for_ocr(image_bytes: bytes):
    from PIL import Image

    try:
        import cv2
        import numpy as np

        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if img is None:
            return Image.open(io.BytesIO(image_bytes))
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        denoised = cv2.fastNlMeansDenoising(enhanced, None, 10, 7, 21)
        binary = cv2.adaptiveThreshold(
            denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
        )
        return Image.fromarray(binary)
    except Exception:
        return Image.open(io.BytesIO(image_bytes))


def _ocr_image_bytes(image_bytes: bytes) -> tuple[str, str]:
    try:
        import pytesseract
    except Exception:
        return "", "ocr_unavailable"
    img = _preprocess_for_ocr(image_bytes)
    try:
        return pytesseract.image_to_string(img), "extracted_text"
    except Exception as exc:
        return "", f"ocr_error:{exc}"


def _ocr_pdf_fallback(data: bytes) -> str:
    try:
        from pdf2image import convert_from_bytes
    except Exception:
        return ""
    try:
        images = convert_from_bytes(data, dpi=200)
    except Exception:
        return ""
    chunks = []
    for img in images:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        text, _ = _ocr_image_bytes(buf.getvalue())
        if text.strip():
            chunks.append(text)
    return "\n--- PAGE BREAK ---\n".join(chunks)


def extract_pdf(data: bytes, page_range: str | None = None) -> IngestResult:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    if page_range:
        pages = _slice_pages(pages, page_range)

    joined = "\n--- PAGE BREAK ---\n".join(pages)
    if not joined.strip():
        ocr_text = _ocr_pdf_fallback(data)
        if ocr_text.strip():
            return IngestResult(
                raw_text=ocr_text,
                data_type="pdf_ocr",
                page_count=len(reader.pages),
                metadata={"ocr": True},
            )
        return IngestResult(
            raw_text="",
            data_type="pdf",
            page_count=len(reader.pages),
            metadata={"ocr_needed": True, "empty_text": True},
        )

    meta: dict = {"page_count": len(reader.pages)}
    if len(reader.pages) > settings.max_pdf_pages_full and not page_range:
        meta["suggest_page_range"] = True
    return IngestResult(
        raw_text=joined,
        data_type="pdf",
        page_count=len(reader.pages),
        metadata=meta,
    )


def extract_excel(data: bytes) -> IngestResult:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    sheet_outputs: list[str] = []
    for sheet in wb.worksheets:
        sheet_data = [f"Sheet Name: {sheet.title}"]
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None and str(cell).strip() != "" for cell in row):
                sheet_data.append(" | ".join("" if c is None else str(c) for c in row))
        sheet_outputs.append("\n".join(sheet_data))
    return IngestResult(
        raw_text="\n\n".join(sheet_outputs),
        data_type="excel_matrix",
        metadata={"sheets": len(wb.worksheets)},
    )


def extract_docx(data: bytes) -> IngestResult:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return IngestResult(raw_text="\n".join(parts), data_type="docx")


def optimize_image(data: bytes, prompt: str) -> IngestResult:
    from PIL import Image

    if OCR_INTENT.search(prompt or ""):
        text, mode = _ocr_image_bytes(data)
        if text.strip():
            return IngestResult(
                raw_text=text, data_type="image_ocr_text", metadata={"mode": mode}
            )

    img = Image.open(io.BytesIO(data))
    max_dim = settings.vision_max_dimension
    if max(img.size) > max_dim:
        img.thumbnail((max_dim, max_dim), Image.Resampling.LANCZOS)
    out = io.BytesIO()
    rgb = img.convert("RGB")
    rgb.save(out, format="JPEG", quality=settings.jpeg_quality, optimize=True)
    b64 = base64.b64encode(out.getvalue()).decode("ascii")
    return IngestResult(
        raw_text="[IMAGE INPUT DETECTED: BOUND VIA MULTIMODAL OVERLAY]",
        data_type="optimized_image_pixels",
        image_b64=b64,
        image_mime="image/jpeg",
        metadata={"mode": "visual", "size": list(rgb.size)},
    )


def extract_text_file(data: bytes, ext: str) -> IngestResult:
    return IngestResult(raw_text=data.decode("utf-8", errors="ignore"), data_type=ext or "text")


def ingest_file(
    filename: str,
    data: bytes,
    prompt: str = "",
    page_range: str | None = None,
) -> IngestResult:
    ext = _ext(filename)
    if ext == "pdf":
        return extract_pdf(data, page_range=page_range)
    if ext in {"xlsx", "xls"}:
        return extract_excel(data)
    if ext == "docx":
        return extract_docx(data)
    if ext == "doc":
        return IngestResult(
            raw_text=data.decode("utf-8", errors="ignore"),
            data_type="doc_legacy",
            metadata={"warning": "legacy .doc best-effort decode; prefer .docx"},
        )
    if ext in {"png", "jpg", "jpeg", "webp", "gif", "bmp"}:
        return optimize_image(data, prompt)
    if ext in {
        "txt", "md", "csv", "json", "tsv", "log", "py", "ts", "js", "toml", "yaml", "yml",
    }:
        return extract_text_file(data, ext)
    return extract_text_file(data, ext or "bin")


def merge_ingests(parts: list[IngestResult]) -> IngestResult:
    if not parts:
        return IngestResult()
    if len(parts) == 1:
        return parts[0]
    texts: list[str] = []
    image = mime = None
    types: list[str] = []
    for i, p in enumerate(parts, 1):
        types.append(p.data_type)
        if p.image_b64 and image is None:
            image, mime = p.image_b64, p.image_mime
        if p.raw_text.strip():
            texts.append(f"### ATTACHMENT {i} [{p.data_type}]\n{p.raw_text}")
    return IngestResult(
        raw_text="\n\n".join(texts),
        data_type="+".join(types),
        image_b64=image,
        image_mime=mime,
        metadata={"attachments": len(parts)},
    )
